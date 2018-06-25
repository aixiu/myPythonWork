#!/usr/bin/env python
# -*- coding: utf-8 -*-

import docx

# 例一
# # 新建一个文档 括号为空即新建，也可写入文件名即为打开该文档
# document = docx.Document()

# document.add_paragraph('Hello,Word!')   # 写入一个段落

# document.save('demo.docx')   #  保存文件


# 例二
document = docx.Document()

paragraph = document.add_paragraph('作者：苏轼', 'Subtitle')

prior_paragraph = paragraph.insert_paragraph_before('水调歌头', 'Title')

document.add_paragraph(
'''
明月几时有，把酒问青天。 
不知天上宫阙，今夕是何年？ 
我欲乘风归去，又恐琼楼玉宇，高处不胜寒。 
起舞弄清影，何似在人间！ 

转朱阁，低绮户，照无眠。 
不应有恨，何事长向别时圆？ 
人有悲欢离合，月有阴晴圆缺，此事古难全。 
但愿人长久，千里共婵娟。
''', 'Body Text 2')

# 标题控制
# document.add_heading('The Real meaning of the nuiverse')
# document.add_heading('Heading 0', level = 0)
# document.add_heading('Heading 1', level = 1)
# document.add_heading('Heading 2', level = 2)

# for p in document.paragraphs:
#     print(len(p.text))
#     print(p.style.name)
# print()

from docx.enum.style import WD_STYLE_TYPE

styles = document.style
print('\n'.join([s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]))

document.save('test.docx')